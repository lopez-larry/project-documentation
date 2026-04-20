# pip install python-docx
# python3 md_to_word.py
#
"""
Create a Word document for every .md file in ./Templates
and write the .docx files to ./Word-Template.

Supported markdown features:
- Headings (# through ######)
- Bullet lists (-, *, +)
- Numbered lists (1. 2. 3.)
- Code fences ```...```
- Blank lines / paragraph breaks
- Basic paragraphs

Notes:
- This is a lightweight converter, not a full Markdown renderer.
- It requires: python-docx
    pip install python-docx
"""

from pathlib import Path
import re
import sys

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    print("Missing dependency: python-docx")
    print("Install it with: pip install python-docx")
    sys.exit(1)


TEMPLATES_DIR = Path("./Templates")
OUTPUT_DIR = Path("./Word-Template")


def ensure_directories() -> None:
    if not TEMPLATES_DIR.exists():
        raise FileNotFoundError(f"Input directory does not exist: {TEMPLATES_DIR.resolve()}")
    if not TEMPLATES_DIR.is_dir():
        raise NotADirectoryError(f"Input path is not a directory: {TEMPLATES_DIR.resolve()}")

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def is_bullet_line(line: str) -> bool:
    return bool(re.match(r"^\s*[-*+]\s+", line))


def is_numbered_line(line: str) -> bool:
    return bool(re.match(r"^\s*\d+\.\s+", line))


def is_heading_line(line: str) -> bool:
    return bool(re.match(r"^\s{0,3}#{1,6}\s+", line))


def extract_heading(line: str) -> tuple[int, str]:
    match = re.match(r"^\s{0,3}(#{1,6})\s+(.*)$", line)
    if not match:
        return 0, line.strip()
    level = len(match.group(1))
    text = match.group(2).strip()
    return level, text


def strip_bullet_marker(line: str) -> str:
    return re.sub(r"^\s*[-*+]\s+", "", line).strip()


def strip_number_marker(line: str) -> str:
    return re.sub(r"^\s*\d+\.\s+", "", line).strip()


def add_code_paragraph(doc: Document, code_line: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(code_line.rstrip("\n"))
    font = run.font
    font.name = "Courier New"
    font.size = Pt(10)


def add_paragraph(doc: Document, text: str) -> None:
    if text.strip():
        doc.add_paragraph(text.strip())
    else:
        doc.add_paragraph("")


def convert_markdown_to_docx(md_path: Path, docx_path: Path) -> None:
    doc = Document()

    # Optional document defaults
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    in_code_block = False

    with md_path.open("r", encoding="utf-8") as f:
        for raw_line in f:
            line = raw_line.rstrip("\n")

            # Toggle fenced code blocks
            if line.strip().startswith("```"):
                in_code_block = not in_code_block
                continue

            if in_code_block:
                add_code_paragraph(doc, line)
                continue

            if not line.strip():
                doc.add_paragraph("")
                continue

            if is_heading_line(line):
                level, text = extract_heading(line)
                # python-docx heading levels are 0-9; use markdown level directly
                doc.add_heading(text, level=level)
                continue

            if is_bullet_line(line):
                text = strip_bullet_marker(line)
                doc.add_paragraph(text, style="List Bullet")
                continue

            if is_numbered_line(line):
                text = strip_number_marker(line)
                doc.add_paragraph(text, style="List Number")
                continue

            add_paragraph(doc, line)

    doc.save(docx_path)


def main() -> None:
    ensure_directories()

    md_files = sorted(TEMPLATES_DIR.glob("*.md"))

    if not md_files:
        print(f"No .md files found in: {TEMPLATES_DIR.resolve()}")
        return

    converted_count = 0

    for md_file in md_files:
        output_file = OUTPUT_DIR / f"{md_file.stem}.docx"
        try:
            convert_markdown_to_docx(md_file, output_file)
            print(f"Created: {output_file}")
            converted_count += 1
        except Exception as exc:
            print(f"Failed to convert {md_file.name}: {exc}")

    print(f"\nDone. Converted {converted_count} file(s).")


if __name__ == "__main__":
    main()